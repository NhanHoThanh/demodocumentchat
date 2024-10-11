import asyncio
from docx import Document as docxDocument
from openai import OpenAI
import requests
from supabase import create_client, Client
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import SupabaseVectorStore
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_openai import OpenAIEmbeddings
from dotenv import load_dotenv
import os
from langchain_core.documents import Document


class ChatBot:
    def __init__(self):
        load_dotenv()
        self.OpenAIClient = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-ada-002", api_key=os.getenv('OPENAI_API_KEY'))

        supabase_url = os.getenv('SUPABASE_URL')
        supabase_key = os.getenv('SUPABASE_KEY')

        self.supabase_client: Client = create_client(
            supabase_url, supabase_key)

        self.vector_store = SupabaseVectorStore(
            client=self.supabase_client,
            embedding=self.embeddings,
            table_name="documents",
            query_name="match_documents",
        )

    def load_and_split_document(self, file_path):
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                document = file.read()
        elif file_path.endswith('.docx'):
            doc = docxDocument(file_path)
            document = '\n'.join([para.text for para in doc.paragraphs])
        else:
            raise ValueError(
                "Unsupported file format. Only .txt and .docx are supported.")

        splitter = RecursiveCharacterTextSplitter()
        chunks = splitter.split_text(document)
        return chunks

    def embed_text(self, chunks):
        embedded_chunks = self.embeddings.embed_documents(chunks)
        return embedded_chunks

    def store_vectors_in_supabase(self, embedded_chunks, text_chunks):

        highest_id = 0

        response = self.supabase_client.table('documents').select(
            'id').order('id', desc=True).limit(1).execute()
        if response.data:
            highest_id = response.data[0]['id']

        ids = list(range(highest_id + 1, highest_id + 1 + len(embedded_chunks)))
        documents = [
            Document(page_content=chunk, metadata={"id": id_})
            for chunk, id_ in zip(text_chunks, ids)
        ]
        self.vector_store.add_vectors(embedded_chunks, documents, ids)

    def process_query(self, question):
        question_vector = self.embeddings.embed_query(question)

        response = self.vector_store.similarity_search_by_vector(
            question_vector, k=5)
        # return response
        relevant_chunks = [doc.page_content for doc in response]

        context = ' '.join(relevant_chunks)
        # return context
        prompt = f"Context: {context}\n\nQuestion: {question}\nAnswer:"
        chat_response = self.OpenAIClient.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "assistant",
                    "content": prompt
                }
            ],
            max_tokens=150
        )

        chatCompletion = chat_response.choices[0].message

        return chatCompletion.content

    # def run(self, file_path, question):
    #     # Load and split the document
    #     chunks = self.load_and_split_document(file_path)

    #     # Embed the text
    #     embedded_chunks = self.embed_text(chunks)

    #     # Store vectors in Supabase
    #     self.store_vectors_in_supabase(embedded_chunks)

    #     # Process the query
    #     answer = self.process_query(question)
    #     print(f"Answer: {answer}")


# async def main():
#     chatbot = ChatBot()
#     file_path = 'anqp.docx'
#     chunks = chatbot.load_and_split_document(file_path)
#     embedded_chunks = chatbot.embed_text(chunks)
#     await chatbot.store_vectors_in_supabase(embedded_chunks, chunks)
    # question = "đưa ra các thực trạng và giải pháp hiện nay"
    # answer = chatbot.process_query(question)
    # print(f"Answer: {answer}")


# asyncio.run(main())

if __name__ == "__main__":
    # main()
    chatbot = ChatBot()
    # file_path = 'anqp.docx'
    # chunks = chatbot.load_and_split_document(file_path)
    # # print(chunks)
    # embedded_chunks = chatbot.embed_text(chunks)
    # chatbot.store_vectors_in_supabase(embedded_chunks, chunks)
    question = "đưa ra các thực trạng và giải pháp hiện nay"
    answer = chatbot.process_query(question)
    print(f"Answer: {answer}")
    # print(embedded_chunks[0])
